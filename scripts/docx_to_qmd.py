"""
Convert a CHA narrative Word document (.docx) into a Quarto (.qmd) chapter file.

The Word document must follow the conventions defined in the CHA Chapter Template:

- Heading 1        -> chapter title (YAML front matter)
- Heading 2 / 3    -> ## / ### section headings
- Normal text      -> body paragraphs with inline formatting preserved
- List Bullet      -> Markdown bulleted list items
- "NOTE: ..." para -> ::: {.callout-note} block
- "SOURCE: ..." para -> ::: {.callout-note collapse="true"} source block
- Superscript runs -> citation anchor HTML  (<a href="#cite-N"><sup>N</sup></a>)
- [Table: id]      -> table cross-reference + render_table_object code block
- [Figure: id]     -> figure cross-reference + render_figure_object code block
- Last Word table  -> parsed as citations table for References section

Usage:
    python scripts/docx_to_qmd.py narrative.docx
    python scripts/docx_to_qmd.py narrative.docx --output chapters/03-demographics.qmd
    python scripts/docx_to_qmd.py narrative.docx --workbook-var CHA_WORKBOOK_PATH
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class Citation:
    number: str
    author: str
    year: str
    title: str
    url: str
    accessed: str


@dataclass
class ConversionState:
    title: str = ""
    body_lines: list[str] = field(default_factory=list)
    citations: list[Citation] = field(default_factory=list)
    table_ids_seen: list[str] = field(default_factory=list)
    figure_ids_seen: list[str] = field(default_factory=list)
    workbook_var: str = "CHA_WORKBOOK_PATH"
    object_render_mode: str = "inline"
    include_dir: str = "_generated/ch04-objects"
    in_bullet_run: bool = False
    in_citations_zone: bool = False


# ---------------------------------------------------------------------------
# Citation table parser
# ---------------------------------------------------------------------------

CITATION_HEADERS = {"number", "author", "organization", "year", "title", "url", "accessed"}


def _is_citations_table(table) -> bool:
    """Detect whether a Word table is the Citations Table by checking header cells."""
    if len(table.rows) < 2 or len(table.columns) < 4:
        return False
    header_texts = {cell.text.strip().lower().split("/")[0].strip() for cell in table.rows[0].cells}
    return len(header_texts & CITATION_HEADERS) >= 3


def _parse_citations_table(table) -> list[Citation]:
    """Extract citation rows from the last table in the document."""
    headers = [cell.text.strip() for cell in table.rows[0].cells]

    def _col_idx(candidates: list[str]) -> int | None:
        for h_idx, h in enumerate(headers):
            if h.lower().replace("/", " ").replace("  ", " ") in [c.lower() for c in candidates]:
                return h_idx
            for c in candidates:
                if c.lower() in h.lower():
                    return h_idx
        return None

    idx_num = _col_idx(["Number", "#"])
    idx_author = _col_idx(["Author", "Author / Organization", "Organization"])
    idx_year = _col_idx(["Year"])
    idx_title = _col_idx(["Title"])
    idx_url = _col_idx(["URL", "Link"])
    idx_accessed = _col_idx(["Accessed"])

    citations: list[Citation] = []
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        num = cells[idx_num] if idx_num is not None else ""
        if not num:
            continue
        citations.append(Citation(
            number=num,
            author=cells[idx_author] if idx_author is not None else "",
            year=cells[idx_year] if idx_year is not None else "",
            title=cells[idx_title] if idx_title is not None else "",
            url=cells[idx_url] if idx_url is not None else "",
            accessed=cells[idx_accessed] if idx_accessed is not None else "",
        ))
    return citations


# ---------------------------------------------------------------------------
# Inline run → Markdown conversion
# ---------------------------------------------------------------------------

_TABLE_REF_RE = re.compile(r"\[Table:\s*([^\]]+)\]", re.IGNORECASE)
_FIGURE_REF_RE = re.compile(r"\[Figure:\s*([^\]]+)\]", re.IGNORECASE)


def _run_to_markdown(run: Run, state: ConversionState) -> str:
    """Convert a single python-docx Run to Markdown/HTML text."""
    text = run.text
    if not text:
        return ""

    if run.font.superscript:
        num = text.strip()
        citation = next((c for c in state.citations if c.number == num), None)
        if citation and citation.url:
            tooltip = f"{citation.author}, {citation.year}, {citation.url}, accessed {citation.accessed}"
            return f'<a href="#cite-{num}" title="{tooltip}"><sup>{num}</sup></a>'
        return f"<sup>{num}</sup>"

    is_bold = run.bold
    is_italic = run.italic

    # Check for table/figure reference patterns inside the run text.
    # These are typed as bold [Table: xyz] or [Figure: xyz].
    def _replace_table_ref(m: re.Match) -> str:
        obj_id = m.group(1).strip()
        state.table_ids_seen.append(obj_id)
        return f"[see @tbl-{obj_id}]"

    def _replace_figure_ref(m: re.Match) -> str:
        obj_id = m.group(1).strip()
        state.figure_ids_seen.append(obj_id)
        return f"[see @fig-{obj_id}]"

    text = _TABLE_REF_RE.sub(_replace_table_ref, text)
    text = _FIGURE_REF_RE.sub(_replace_figure_ref, text)

    # After reference substitution, skip wrapping the reference text itself in
    # bold markers -- the [see @...] syntax should not be bolded.
    has_ref = "[see @" in text

    if is_bold and is_italic and not has_ref:
        return f"***{text}***"
    if is_bold and not has_ref:
        return f"**{text}**"
    if is_italic:
        return f"*{text}*"
    return text


def _paragraph_to_markdown(para: Paragraph, state: ConversionState) -> str:
    """Convert all runs in a paragraph into a single Markdown string."""
    parts: list[str] = []
    for run in para.runs:
        parts.append(_run_to_markdown(run, state))
    return "".join(parts)


# ---------------------------------------------------------------------------
# Instructions page detection
# ---------------------------------------------------------------------------

def _is_instructions_heading(para: Paragraph) -> bool:
    return (
        para.style.name.startswith("Heading")
        and "INSTRUCTIONS" in para.text.upper()
        and "DELETE" in para.text.upper()
    )


# ---------------------------------------------------------------------------
# Code block generators
# ---------------------------------------------------------------------------

def _table_code_block(obj_id: str, workbook_var: str) -> str:
    label = f"tbl-{obj_id}"
    return (
        f"```{{python}}\n"
        f"#| echo: false\n"
        f"#| warning: false\n"
        f"#| message: false\n"
        f"#| label: {label}\n"
        f'#| tbl-cap: ""\n'
        f"from scripts.cha_registry_renderer import render_table_object\n"
        f"\n"
        f"render_table_object(\n"
        f'    object_id="{label}",\n'
        f"    workbook_path={workbook_var},\n"
        f")\n"
        f"```"
    )


def _figure_code_block(obj_id: str, workbook_var: str) -> str:
    label = f"fig-{obj_id}"
    return (
        f"```{{python}}\n"
        f"#| echo: false\n"
        f"#| warning: false\n"
        f"#| message: false\n"
        f"#| label: {label}\n"
        f'#| fig-cap: ""\n'
        f"from scripts.cha_registry_renderer import render_figure_object\n"
        f"\n"
        f'render_figure_object(figure_id="{label}", workbook_path={workbook_var}).show()\n'
        f"```"
    )


def _object_include_shortcode(obj_id: str, include_dir: str) -> str:
    return f"{{{{< include {include_dir}/{obj_id}.qmd >}}}}"


# ---------------------------------------------------------------------------
# Workbook path preamble
# ---------------------------------------------------------------------------

def _workbook_preamble(workbook_var: str) -> str:
    return (
        f"```{{python}}\n"
        f"#| echo: false\n"
        f"#| warning: false\n"
        f"#| message: false\n"
        f"import os\n"
        f"import sys\n"
        f"from pathlib import Path\n"
        f"\n"
        f"cwd = Path.cwd()\n"
        f'project_root = cwd if (cwd / "scripts").exists() else cwd.parent\n'
        f"if str(project_root) not in sys.path:\n"
        f"    sys.path.insert(0, str(project_root))\n"
        f"\n"
        f'{workbook_var} = project_root / "data" / "raw" / "Mid-Hudson Regional Community Health Assessment 2025 Data File.xlsx"\n'
        f"```"
    )


# ---------------------------------------------------------------------------
# References section builder
# ---------------------------------------------------------------------------

def _build_references_section(citations: list[Citation]) -> str:
    if not citations:
        return ""
    lines = [
        "",
        "## References",
        "",
        "<details>",
        '<summary><strong>Click to view citations</strong></summary>',
        "",
    ]
    for c in citations:
        url_part = f"<{c.url}>" if c.url else ""
        accessed_part = f", accessed {c.accessed}" if c.accessed else ""
        lines.append(
            f'<span id="cite-{c.number}">{c.number}. {c.author}, {c.year}, '
            f"{url_part}{accessed_part}</span>"
        )
        lines.append("")
    lines.append("</details>")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Callout content extractor
# ---------------------------------------------------------------------------

def _extract_callout_content(para: Paragraph, prefix: str, state: ConversionState) -> str:
    """Extract the content after a NOTE:/SOURCE: prefix, skipping the prefix runs.

    Walks the runs character-by-character to consume exactly the prefix text,
    then renders remaining runs normally. This avoids stray bold/italic markers
    that occur when the prefix sits in its own bold run.
    """
    prefix_stripped = prefix.strip()
    chars_remaining = len(prefix_stripped)
    content_parts: list[str] = []

    for run in para.runs:
        run_text = run.text
        if chars_remaining > 0:
            # This run is still part of the prefix — consume characters
            text_for_prefix = run_text.lstrip() if chars_remaining == len(prefix_stripped) else run_text
            if chars_remaining >= len(text_for_prefix):
                chars_remaining -= len(text_for_prefix)
                continue
            # Partially consumed: take the leftover after prefix chars
            leftover = text_for_prefix[chars_remaining:].lstrip()
            chars_remaining = 0
            if not leftover:
                continue
            # Render the leftover with the same formatting as this run
            if run.italic:
                content_parts.append(f"*{leftover}*")
            elif run.bold:
                content_parts.append(f"**{leftover}**")
            else:
                content_parts.append(leftover)
            continue
        content_parts.append(_run_to_markdown(run, state))

    return "".join(content_parts)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert_docx_to_qmd(
    docx_path: str | Path,
    workbook_var: str = "CHA_WORKBOOK_PATH",
    object_render_mode: str = "inline",
    include_dir: str = "_generated/ch04-objects",
) -> str:
    """Read a .docx file and return the generated .qmd content as a string."""
    doc = Document(str(docx_path))

    state = ConversionState(
        workbook_var=workbook_var,
        object_render_mode=object_render_mode,
        include_dir=include_dir,
    )

    # Parse citations from the last table that matches the schema
    for table in reversed(doc.tables):
        if _is_citations_table(table):
            state.citations = _parse_citations_table(table)
            break

    # Walk paragraphs, skipping the instructions page
    skip_instructions = True
    past_first_heading1 = False

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()

        # Skip the instructions page: everything from the INSTRUCTIONS heading
        # until we hit the first non-instructions Heading 1.
        if skip_instructions:
            if _is_instructions_heading(para):
                continue
            if style.startswith("Heading 1") and not _is_instructions_heading(para):
                skip_instructions = False
            else:
                continue

        # Empty paragraph -> blank line (but avoid stacking multiple blanks)
        if not text:
            if state.body_lines and state.body_lines[-1] != "":
                state.body_lines.append("")
            state.in_bullet_run = False
            continue

        # Once we hit the Citations heading, skip all remaining normal paragraphs
        # (they are citation table instructions, not chapter content).
        if state.in_citations_zone:
            if style.startswith("Heading"):
                pass  # fall through to heading handlers below
            else:
                continue

        # Heading 1 -> chapter title (first one only)
        if style.startswith("Heading 1"):
            if not past_first_heading1:
                state.title = text
                past_first_heading1 = True
                continue
            # Subsequent Heading 1s (like "Citations") are section headers
            if text.lower() in ("citations", "references"):
                continue
            state.body_lines.append(f"## {text}")
            state.body_lines.append("")
            state.in_bullet_run = False
            continue

        # Heading 2/3
        if style.startswith("Heading 2"):
            if text.lower() in ("citations", "references"):
                state.in_citations_zone = True
                continue
            state.in_citations_zone = False
            state.in_bullet_run = False
            state.body_lines.append(f"## {text}")
            state.body_lines.append("")
            continue

        if style.startswith("Heading 3"):
            state.in_bullet_run = False
            state.body_lines.append(f"### {text}")
            state.body_lines.append("")
            continue

        if style.startswith("Heading 4"):
            state.in_bullet_run = False
            state.body_lines.append(f"#### {text}")
            state.body_lines.append("")
            continue

        # Citations heading paragraph (sometimes typed as normal text)
        if text.lower() in ("citations", "references") and not para.runs:
            continue

        # NOTE: callout
        if text.upper().startswith("NOTE:"):
            state.in_bullet_run = False
            note_content = _extract_callout_content(para, "NOTE:", state)
            state.body_lines.append("::: {.callout-note}")
            state.body_lines.append("## Note")
            state.body_lines.append("")
            state.body_lines.append(note_content)
            state.body_lines.append(":::")
            state.body_lines.append("")
            continue

        # SOURCE: callout
        if text.upper().startswith("SOURCE:"):
            state.in_bullet_run = False
            src_content = _extract_callout_content(para, "SOURCE:", state)
            state.body_lines.append('::: {.callout-note collapse="true"}')
            state.body_lines.append("## Source")
            state.body_lines.append("")
            state.body_lines.append(src_content)
            state.body_lines.append(":::")
            state.body_lines.append("")
            continue

        # Bulleted list
        if "List Bullet" in style or "List Paragraph" in style:
            md_text = _paragraph_to_markdown(para, state)
            state.body_lines.append(f"-   {md_text}")
            state.in_bullet_run = True
            continue

        # Normal paragraph — add blank line after bullet run for proper spacing
        if state.in_bullet_run:
            state.body_lines.append("")
        state.in_bullet_run = False

        # Track table/figure refs found before rendering so we can insert code blocks after
        tables_before = len(state.table_ids_seen)
        figures_before = len(state.figure_ids_seen)

        md_text = _paragraph_to_markdown(para, state)

        state.body_lines.append(md_text)
        state.body_lines.append("")

        new_tables = state.table_ids_seen[tables_before:]
        new_figures = state.figure_ids_seen[figures_before:]

        if state.object_render_mode == "include":
            include_ids: list[str] = []
            seen: set[str] = set()
            for obj_id in new_figures + new_tables:
                if obj_id not in seen:
                    include_ids.append(obj_id)
                    seen.add(obj_id)
            for obj_id in include_ids:
                state.body_lines.append(_object_include_shortcode(obj_id, state.include_dir))
                state.body_lines.append("")
        else:
            for tid in new_tables:
                state.body_lines.append(_table_code_block(tid, state.workbook_var))
                state.body_lines.append("")

            for fid in new_figures:
                state.body_lines.append(_figure_code_block(fid, state.workbook_var))
                state.body_lines.append("")

    # Build YAML front matter
    yaml_lines = [
        "---",
        f'title: "{state.title}"',
        "jupyter: python3",
        "toc: true",
        "toc-depth: 3",
        "execute:",
        "  freeze: auto",
        "---",
    ]

    # Assemble final output
    parts = [
        "\n".join(yaml_lines),
        "",
        _workbook_preamble(state.workbook_var),
        "",
    ]

    # Body
    body = "\n".join(state.body_lines).strip()
    parts.append(body)

    # References
    refs = _build_references_section(state.citations)
    if refs:
        parts.append(refs)

    return "\n".join(parts) + "\n"


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(
        description="Convert a CHA narrative Word document to Quarto .qmd format."
    )
    parser.add_argument("docx", help="Path to the input .docx file")
    parser.add_argument(
        "--output", "-o",
        default="",
        help="Output .qmd file path (default: same name as input with .qmd extension)",
    )
    parser.add_argument(
        "--workbook-var",
        default="CHA_WORKBOOK_PATH",
        help="Python variable name for the workbook path (default: CHA_WORKBOOK_PATH)",
    )
    parser.add_argument(
        "--object-render-mode",
        choices=["inline", "include"],
        default="inline",
        help="Render [Table:/Figure:] references as inline python blocks or include shortcodes.",
    )
    parser.add_argument(
        "--include-dir",
        default="_generated/ch04-objects",
        help="Include directory used when --object-render-mode=include.",
    )
    args = parser.parse_args(argv)

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output) if args.output else docx_path.with_suffix(".qmd")

    qmd_content = convert_docx_to_qmd(
        docx_path,
        workbook_var=args.workbook_var,
        object_render_mode=args.object_render_mode,
        include_dir=args.include_dir,
    )

    output_path.write_text(qmd_content, encoding="utf-8")
    print(f"Converted: {docx_path} -> {output_path}")


if __name__ == "__main__":
    main()
