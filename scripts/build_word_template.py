"""
Build the CHA Chapter Word template programmatically using python-docx.

Run once:  python scripts/build_word_template.py
Output:    templates/CHA_Chapter_Template.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "templates" / "CHA_Chapter_Template.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_run(paragraph, text: str, *, bold=False, italic=False, superscript=False,
             size: Pt | None = None, color: RGBColor | None = None, font_name: str | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run


def _set_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(shd)


def _style_header_cell(cell, text: str):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_shading(cell, "2F5496")


def _style_body_cell(cell, text: str, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold


# ---------------------------------------------------------------------------
# Page break helper
# ---------------------------------------------------------------------------

def _add_page_break(doc: Document):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Instructions page
# ---------------------------------------------------------------------------

def _build_instructions_page(doc: Document):
    title = doc.add_heading("INSTRUCTIONS – DELETE THIS PAGE BEFORE SUBMITTING", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph(
        "This template lets you write CHA chapter narrative in Microsoft Word. "
        "A converter script will transform this document into Quarto (.qmd) format "
        "automatically. Follow these conventions so the converter can parse your content."
    )

    doc.add_heading("Headings", level=2)
    doc.add_paragraph(
        'Use Word\'s built-in Heading 2 style for major sections (e.g. "Demographic '
        'Characteristics") and Heading 3 for subsections (e.g. "Population"). '
        "Do NOT use Heading 1 except for the chapter title on the next page."
    )

    doc.add_heading("Table and Figure References", level=2)
    p = doc.add_paragraph()
    _add_run(p, "To reference a data table, type: ")
    _add_run(p, "[Table: population-demographics]", bold=True)
    p2 = doc.add_paragraph()
    _add_run(p2, "To reference a figure/chart, type: ")
    _add_run(p2, "[Figure: labor-force]", bold=True)
    doc.add_paragraph(
        "The ID after the colon must match the object_id in the Excel workbook "
        "(without the tbl- or fig- prefix). The converter adds the prefix automatically."
    )

    doc.add_heading("Citations (Superscript Numbers)", level=2)
    doc.add_paragraph(
        "To cite a source inline, type the citation number and format it as "
        "superscript using Word's superscript button (or Ctrl+Shift+=). "
        "For example, to cite source #7, type 7 and make it superscript."
    )
    p3 = doc.add_paragraph("Example: The population grew 4.7% from 2010 to 2020.")
    _add_run(p3, "7", superscript=True)
    doc.add_paragraph(
        "Then fill in the citation details in the Citations Table at the end of "
        "this document. The converter will build the proper HTML anchors and the "
        "References section automatically."
    )

    doc.add_heading("Notes and Callouts", level=2)
    doc.add_paragraph(
        'To create a note/callout box, start a paragraph with the word "NOTE:" '
        "in bold. Everything after NOTE: becomes the callout content."
    )
    p4 = doc.add_paragraph()
    _add_run(p4, "NOTE: ", bold=True)
    _add_run(p4, "The American Community Survey includes a question that intends to "
                  "capture current sex; there are no questions about gender.")

    doc.add_heading("Source Callouts", level=2)
    doc.add_paragraph(
        'To create a collapsible source box, start a paragraph with "SOURCE:" '
        "in bold, then provide the full source text."
    )
    p5 = doc.add_paragraph()
    _add_run(p5, "SOURCE: ", bold=True)
    _add_run(p5, "US Census Bureau; American Community Survey, 2023 American Community "
                  "Survey 5-Year Estimates, Table S0101, April 2025")

    doc.add_heading("Bulleted Lists", level=2)
    doc.add_paragraph(
        "Use Word's standard bulleted list. Each bullet becomes a Markdown list item."
    )
    doc.add_paragraph("First bullet item", style="List Bullet")
    doc.add_paragraph("Second bullet item", style="List Bullet")

    doc.add_heading("Bold and Italic Text", level=2)
    p6 = doc.add_paragraph()
    _add_run(p6, "Bold text: ", bold=True)
    _add_run(p6, "becomes **bold** in Markdown. ")
    _add_run(p6, "Italic text: ", italic=True)
    _add_run(p6, "becomes *italic* in Markdown.")

    doc.add_heading("What NOT To Do", level=2)
    doc.add_paragraph("Do NOT paste data tables into this document. "
                       "Tables are rendered from the Excel workbook automatically.")
    doc.add_paragraph("Do NOT use Heading 1 for anything other than the chapter title.")
    doc.add_paragraph("Do NOT change or delete the Citations Table structure at the end.")

    _add_page_break(doc)


# ---------------------------------------------------------------------------
# Example body content
# ---------------------------------------------------------------------------

def _build_example_body(doc: Document):
    # Chapter title (Heading 1)
    doc.add_heading("Chapter Title Goes Here", level=1)

    # Section heading (Heading 2)
    doc.add_heading("Section Heading", level=2)

    # Subsection (Heading 3)
    doc.add_heading("Subsection Heading", level=3)

    # Normal paragraph with a table reference
    p = doc.add_paragraph(
        "Write your narrative paragraph text here. You can reference data tables "
        "by typing "
    )
    _add_run(p, "[Table: population-demographics]", bold=True)
    _add_run(p, " and the converter will create the proper cross-reference and insert "
                "the code block to render the table from the workbook.")

    # Paragraph with an inline citation superscript
    p2 = doc.add_paragraph(
        "You can cite sources by typing the citation number as superscript. "
        "For example, this sentence has a citation."
    )
    _add_run(p2, "1", superscript=True)

    # A second paragraph with figure reference
    p3 = doc.add_paragraph("Figures work the same way: ")
    _add_run(p3, "[Figure: labor-force]", bold=True)
    _add_run(p3, " will insert the chart from the workbook.")

    # NOTE callout example
    p4 = doc.add_paragraph()
    _add_run(p4, "NOTE: ", bold=True)
    _add_run(p4, "This is an example note. The converter wraps this in a callout box. "
                  "Use this for methodology notes, definitions, or caveats.")

    # SOURCE callout example
    p5 = doc.add_paragraph()
    _add_run(p5, "SOURCE: ", bold=True)
    _add_run(p5, "US Census Bureau; American Community Survey, 2023 American Community "
                  "Survey 5-Year Estimates, Table S0101, April 2025")

    # Bulleted list
    doc.add_paragraph(
        "Here is an example of a bulleted list:"
    )
    doc.add_paragraph("First list item – describe something here", style="List Bullet")
    doc.add_paragraph("Second list item – another description", style="List Bullet")
    doc.add_paragraph("Third list item – one more for good measure", style="List Bullet")

    # Paragraph with italic
    p6 = doc.add_paragraph("According to the ")
    _add_run(p6, "American Community Survey", italic=True)
    _add_run(p6, ", the base population for this category was people aged five years and older.")


# ---------------------------------------------------------------------------
# Citations table
# ---------------------------------------------------------------------------

def _build_citations_table(doc: Document):
    doc.add_heading("Citations", level=2)
    doc.add_paragraph(
        "Fill in one row per source. The Number column must match the superscript "
        "numbers used in the body text above."
    )

    headers = ["Number", "Author / Organization", "Year", "Title", "URL", "Accessed"]
    example_rows = [
        [
            "1",
            "Example Organization",
            "2025",
            "Example Report Title",
            "https://example.com/report",
            "August 2025",
        ],
        ["", "", "", "", "", ""],
        ["", "", "", "", "", ""],
    ]

    table = doc.add_table(rows=1 + len(example_rows), cols=len(headers))
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        _style_header_cell(table.rows[0].cells[i], header)

    for row_idx, row_data in enumerate(example_rows):
        for col_idx, cell_val in enumerate(row_data):
            _style_body_cell(
                table.rows[row_idx + 1].cells[col_idx],
                cell_val,
                bold=(col_idx == 0 and cell_val != ""),
            )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    _build_instructions_page(doc)
    _build_example_body(doc)
    _build_citations_table(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Template saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
